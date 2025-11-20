from flask import (
    Flask,
    render_template,
    request,
    send_file,
    session,
)
from flask_session import Session
from googlesearch import search
import requests
from bs4 import BeautifulSoup
import re
from urllib.parse import quote
import time
import os
import socket
import ssl
import json
from datetime import datetime
from groq import Groq
from concurrent.futures import ThreadPoolExecutor, as_completed
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document
import textwrap
import uuid
import warnings
warnings.filterwarnings("ignore", category=RuntimeWarning)


# =====================================================================================
# CONFIG
# =====================================================================================

report_cache = {}  # temporary in-memory store



NVD_API_KEY = os.getenv("NVD_API_KEY")
if not NVD_API_KEY:
    raise ValueError("NVD_API_KEY environment variable is required")

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable is required")

client = Groq(api_key=GROQ_API_KEY)

MODEL = "moonshotai/kimi-k2-instruct"
DEFAULT_LOOKBACK_DAYS = int(os.getenv("DEFAULT_LOOKBACK_DAYS", 180))

COMMON_TLDS = [".com", ".org", ".net", ".ai", ".dev", ".io"]
RISKY_PORTS = [21, 23, 25, 110, 143, 3389]

# Threat intel cache (CISA KEV)
kev_cache = None
kev_last_fetch = 0
KEV_TTL_SECONDS = 3600  # 1 hour

# =====================================================================================
# FLASK APP INIT
# =====================================================================================

app = Flask(__name__)
app.secret_key = "kejbsughuw"

app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_PERMANENT"] = False
Session(app)

# =====================================================================================
# NVD Vulnerability Search
# =====================================================================================


def search_vulnerabilities_nvd(app_name):
    try:
        q = app_name.strip()
        if q.lower() in {"vs code", "vscode"}:
            q = "Visual Studio Code"
        elif q.lower() in {"chrome", "google chrome"}:
            q = "Google Chrome"
        elif q.lower() in {"firefox", "mozilla firefox"}:
            q = "Mozilla Firefox"

        url = (
            "https://services.nvd.nist.gov/rest/json/cves/2.0"
            f"?keywordSearch={quote(q)}&resultsPerPage=200"
        )
        headers = {"User-Agent": "Mozilla/5.0"}
        if NVD_API_KEY:
            headers["apiKey"] = NVD_API_KEY

        # Be polite to NVD
        time.sleep(0.6)

        response = requests.get(url, headers=headers, timeout=30)
        if response.status_code != 200:
            return []

        data = response.json()
        cves = data.get("vulnerabilities", [])
        parsed = []

        for cve_item in cves:
            try:
                cve_data = cve_item.get("cve", {})
                cve_id = cve_data.get("id", "Unknown")
                descriptions = cve_data.get("descriptions", [])
                description = next(
                    (d.get("value") for d in descriptions if d.get("lang") == "en"),
                    "No description",
                )

                metrics = cve_data.get("metrics", {})
                cvss_score = "N/A"
                for key in ["cvssMetricV31", "cvssMetricV30", "cvssMetricV2"]:
                    if key in metrics and len(metrics[key]) > 0:
                        cvss_score = (
                            metrics[key][0]
                            .get("cvssData", {})
                            .get("baseScore", "N/A")
                        )
                        break

                if isinstance(cvss_score, (int, float)):
                    cvss_score = float(cvss_score)
                elif isinstance(cvss_score, str) and cvss_score.replace(
                    ".", "", 1
                ).isdigit():
                    cvss_score = float(cvss_score)

                parsed.append(
                    {
                        "id": cve_id,
                        "summary": description[:200] + "..."
                        if len(description) > 200
                        else description,
                        "cvss": cvss_score,
                    }
                )
            except Exception:
                continue

        return parsed
    except Exception:
        return []


def analyze_risk(cve_list):
    if not cve_list:
        return "None", True, {}

    critical = sum(
        1
        for cve in cve_list
        if isinstance(cve["cvss"], (int, float)) and 9.0 <= cve["cvss"] <= 10.0
    )
    high = sum(
        1
        for cve in cve_list
        if isinstance(cve["cvss"], (int, float)) and 7.0 <= cve["cvss"] < 9.0
    )
    medium = sum(
        1
        for cve in cve_list
        if isinstance(cve["cvss"], (int, float)) and 4.0 <= cve["cvss"] < 7.0
    )
    low = sum(
        1
        for cve in cve_list
        if isinstance(cve["cvss"], (int, float)) and 0.1 <= cve["cvss"] < 4.0
    )
    none_score = sum(1 for cve in cve_list if cve["cvss"] == 0.0)

    scores = [
        cve["cvss"] for cve in cve_list if isinstance(cve["cvss"], (int, float))
    ]
    if not scores:
        return "None", True, {}

    avg_score = round(sum(scores) / len(scores), 2)
    severity, safe_to_install = "Unknown", False
    if avg_score >= 9.0:
        severity = "Critical"
    elif avg_score >= 8.5:
        severity = "High"
    elif avg_score >= 4.0:
        severity, safe_to_install = "Medium", True
    elif avg_score >= 0.1:
        severity, safe_to_install = "Low", True
    elif avg_score == 0.0:
        severity, safe_to_install = "None", True

    breakdown = {
        "critical": critical,
        "high": high,
        "medium": medium,
        "low": low,
        "none": none_score,
        "total": len(cve_list),
    }
    return severity, safe_to_install, breakdown


# =====================================================================================
# Version & Release Date
# =====================================================================================


def extract_version_and_date(url):
    try:
        response = requests.get(
            url, timeout=15, headers={"User-Agent": "Mozilla/5.0"}
        )
        soup = BeautifulSoup(response.text, "html.parser")
        version, date = None, None

        if "github.com" in url and "/releases" in url:
            version_tag = soup.select_one(
                "a[href*='/tag/'] .css-truncate-target, "
                "h1.d-inline.mr-3, "
                "span.css-truncate-target, "
                "a.Link--primary[href*='/tag/']"
            )
            if version_tag:
                version = version_tag.get_text(strip=True)

            date_tag = soup.select_one("relative-time, time-ago")
            if date_tag:
                date = (
                    date_tag.get("datetime")
                    or date_tag.get("title")
                    or date_tag.get_text(strip=True)
                )
        else:
            text = soup.get_text(" ", strip=True)
            version_match = re.search(r"\bv?(\d+\.\d+\.\d+(?:\.\d+)?)\b", text)
            if version_match:
                version = version_match.group(1)

        return version, date
    except Exception:
        return None, None


# =====================================================================================
# App Description
# =====================================================================================


def fetch_app_description(app_name, search_results):
    # Try Wikipedia with multiple candidates
    try:
        title = app_name.strip()
        candidates = [
            title,
            title.title(),
            title.replace(" ", "_"),
            title.replace(" ", "_").title(),
            f"{title} (software)",
        ]

        if title.lower() in {"vs code", "vscode"}:
            candidates.insert(0, "Visual Studio Code")
        elif title.lower() in {"chrome", "google chrome"}:
            candidates.insert(0, "Google Chrome")

        for cand in candidates:
            try:
                url = (
                    "https://en.wikipedia.org/api/rest_v1/page/summary/"
                    f"{quote(cand)}"
                )
                response = requests.get(
                    url, timeout=10, headers={"User-Agent": "Mozilla/5.0"}
                )
                if response.status_code == 200:
                    data = response.json()
                    extract = data.get("extract")
                    page = (
                        data.get("content_urls", {})
                        .get("desktop", {})
                        .get("page")
                    )
                    if extract and len(extract.strip()) > 10:
                        app.logger.info(
                            f"Found Wikipedia description for: {cand}"
                        )
                        return extract, page
            except Exception:
                continue

    except Exception:
        app.logger.exception("Wikipedia lookup failed")

    # Fallback: GitHub repo description
    try:
        for url in search_results:
            if "github.com" in url and "/releases" not in url:
                r = requests.get(
                    url, timeout=10, headers={"User-Agent": "Mozilla/5.0"}
                )
                soup = BeautifulSoup(r.text, "html.parser")

                selectors = [
                    'meta[property="og:description"]',
                    'meta[name="description"]',
                    ".f4.my-3",
                    ".BorderGrid-cell .f4",
                ]

                for selector in selectors:
                    element = soup.select_one(selector)
                    if element:
                        content = (
                            element.get("content")
                            if element.name == "meta"
                            else element.get_text(strip=True)
                        )
                        if content and len(content.strip()) > 10:
                            app.logger.info(
                                f"Found GitHub description: {content[:100]}..."
                            )
                            return content, url

    except Exception:
        app.logger.exception("GitHub description fallback failed")

    return None, None


# =====================================================================================
# Domain Resolution & Metadata via LLM
# =====================================================================================


def llm_resolve_domain(app_name: str) -> str:
    try:
        prompt = f"Give ONLY the official website domain for '{app_name}'"
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=20,
        )
        candidate = resp.choices[0].message.content.strip().lower()
        if candidate.startswith("http"):
            candidate = candidate.split("//")[-1]
        return candidate.split("/")[0]
    except Exception:
        return ""


def validate_domain(domain: str) -> bool:
    try:
        socket.gethostbyname(domain)
        return True
    except Exception:
        return False


def resolve_application_domain(app_name: str) -> str:
    domain = llm_resolve_domain(app_name)
    if domain and validate_domain(domain):
        return domain

    key = app_name.replace(" ", "").lower()
    for tld in COMMON_TLDS:
        test_domain = key + tld
        if validate_domain(test_domain):
            return test_domain

    return "unknown"


def fetch_app_metadata(app_name: str, domain: str):
    prompt = f"""
    Provide JSON with these keys for '{app_name}' ({domain}):
    - description
    - license
    - app_type
    - latest_version
    - release_date
    - support_status
    - official_link
    """
    try:
        resp = client.chat.completions.create(
            model=MODEL,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=250,
        )
        text = resp.choices[0].message.content.strip()

        if "```json" in text:
            text = text.replace("```json", "").replace("```", "").strip()

        parsed = json.loads(text)

        if isinstance(parsed.get("description"), str):
            try:
                inner = json.loads(parsed["description"])
                parsed.update(inner)
            except Exception:
                pass

        for field in [
            "license",
            "app_type",
            "latest_version",
            "release_date",
            "support_status",
            "official_link",
        ]:
            parsed.setdefault(field, "unknown")

        return parsed

    except Exception as e:
        return {"error": str(e)}


# =====================================================================================
# Security & Compliance (Network Probes, SSL, LLM summary)
# =====================================================================================


def probe_ports(domain: str, ports=None):
    if ports is None:
        ports = [
            21,
            22,
            25,
            53,
            80,
            110,
            143,
            443,
            465,
            587,
            993,
            995,
            3306,
            3389,
            8080,
        ]
    results = {}

    def check_port(port):
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(0.5)
        try:
            if sock.connect_ex((domain, port)) == 0:
                return port
        finally:
            sock.close()

    with ThreadPoolExecutor(max_workers=30) as ex:
        for fut in as_completed(ex.submit(check_port, p) for p in ports):
            port = fut.result()
            if port:
                results[port] = "open"
    return results


def check_ssl(domain: str) -> bool:
    try:
        ctx = ssl.create_default_context()
        with ctx.wrap_socket(socket.socket(), server_hostname=domain) as s:
            s.settimeout(3)
            s.connect((domain, 443))
            return bool(s.getpeercert())
    except Exception:
        return False


def llm_report(app_name, domain, summary, controls):
    prompt = (
        f"Application: {app_name}\n"
        f"Domain: {domain}\n"
        f"Security Summary: {json.dumps(summary)}\n"
        f"Compliance Controls: {json.dumps(controls)}"
    )
    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=400,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"[LLM error: {e}]"


# =====================================================================================
# Threat Intelligence Module
# =====================================================================================


def load_cisa_kev():
    """
    Load CISA Known Exploited Vulnerabilities catalog and cache for KEV_TTL_SECONDS.
    """
    global kev_cache, kev_last_fetch
    now = time.time()
    if kev_cache is not None and (now - kev_last_fetch) < KEV_TTL_SECONDS:
        return kev_cache

    url = (
        "https://www.cisa.gov/sites/default/files/feeds/"
        "known_exploited_vulnerabilities.json"
    )
    try:
        resp = requests.get(url, timeout=15)
        if resp.status_code != 200:
            return []
        data = resp.json()
        kev_cache = data.get("vulnerabilities", [])
        kev_last_fetch = now
        return kev_cache
    except Exception:
        return []


def check_exploitdb_for_cve(cve_id: str) -> bool:
    """
    Quick check: if Exploit-DB search page for this CVE has results.
    """
    try:
        url = f"https://www.exploit-db.com/search?cve={cve_id}"
        resp = requests.get(
            url, timeout=10, headers={"User-Agent": "Mozilla/5.0"}
        )
        # crude check – could be made stricter by parsing html
        return resp.status_code == 200 and "No results" not in resp.text
    except Exception:
        return False


def fetch_cve_patch_references(cve_id: str):
    """
    Use NVD per-CVE endpoint to pull patch/advisory references.
    """
    try:
        base_url = "https://services.nvd.nist.gov/rest/json/cve/2.0"
        params = {"cveId": cve_id}
        headers = {"User-Agent": "Mozilla/5.0"}
        if NVD_API_KEY:
            headers["apiKey"] = NVD_API_KEY

        time.sleep(0.6)
        resp = requests.get(
            base_url, params=params, headers=headers, timeout=20
        )
        if resp.status_code != 200:
            return []

        data = resp.json()
        vulns = data.get("vulnerabilities", [])
        if not vulns:
            return []

        cve = vulns[0].get("cve", {})
        refs = cve.get("references", [])
        patch_refs = []
        for ref in refs:
            url = ref.get("url")
            tags = ref.get("tags", []) or []
            tag_str = " ".join(tags).lower()
            if any(
                kw in tag_str
                for kw in [
                    "patch",
                    "vendor advisory",
                    "third party advisory",
                    "update",
                    "security advisory",
                ]
            ):
                patch_refs.append({"url": url, "tags": tags})
        return patch_refs
    except Exception:
        return []


def search_threat_campaigns(app_name: str, limit: int = 5):
    """
    Google search OSINT for campaigns / active exploitation mentions.
    """
    query = (
        f'"{app_name}" active exploitation OR "ransomware campaign" '
        f'OR "APT" OR "0-day"'
    )
    urls = []
    try:
        for url in search(query, num_results=limit):
            urls.append(url)
    except Exception:
        pass
    return urls

def fetch_recent_attacks(app_name: str, domain: str = None, limit: int = 10):
    """
    Search OSINT (Google News + security sites) for recent security attacks
    related to an application or website. Safe OSINT only.
    """
    # from duckduckgo_search import DDGS
    from ddgs import DDGS

    ddgs = DDGS()

    keywords = [
        f"{app_name} security breach",
        f"{app_name} hacked",
        f"{app_name} cyber attack",
        f"{app_name} data leak",
        f"{app_name} zero-day",
        f"{app_name} exploit",
    ]

    if domain:
        keywords.append(f"{domain} hacked")
        keywords.append(f"{domain} breach")

    results = []

    for query in keywords:
        try:
            news_results = ddgs.news(query, max_results=limit)
            if news_results:
                for item in news_results:
                    results.append({
                        "title": item.get("title"),
                        "url": item.get("url"),
                        "date": item.get("date"),
                        "source": item.get("source")
                    })
        except Exception:
            continue

    # Remove duplicates by URL
    unique = {}
    for r in results:
        unique[r["url"]] = r
    return list(unique.values())



def fetch_threat_intel(app_name: str, domain: str, cve_list):
    """
    Aggregate:
    - CISA KEV matches
    - Exploit-DB presence
    - NVD patch/advisory references
    - Open web campaign links
    """
    cve_ids = [c["id"] for c in (cve_list or []) if c.get("id")]

    # 1) CISA KEV
    kev_data = load_cisa_kev()
    exploited_cves = []
    if kev_data and cve_ids:
        kev_index = {item.get("cveID"): item for item in kev_data}
        for cid in cve_ids:
            if cid in kev_index:
                exploited_cves.append(kev_index[cid])

    # 2) Exploit-DB
    exploits = []
    for cid in cve_ids:
        try:
            if check_exploitdb_for_cve(cid):
                exploits.append(
                    {
                        "cve_id": cid,
                        "source": "Exploit-DB",
                        "url": f"https://www.exploit-db.com/search?cve={cid}",
                    }
                )
        except Exception:
            continue

    # 3) Patch references (cap at 20 CVEs)
    patches = []
    for cid in cve_ids[:20]:
        refs = fetch_cve_patch_references(cid)
        for ref in refs:
            patches.append(
                {
                    "cve_id": cid,
                    "url": ref.get("url"),
                    "tags": ref.get("tags", []),
                    "source": "NVD Reference",
                }
            )

    # 4) Campaign links
    campaigns = search_threat_campaigns(app_name)

    recent_attacks = fetch_recent_attacks(app_name, domain)

    return {
    "exploited_cves": exploited_cves,
    "exploits": exploits,
    "patches": patches,
    "campaigns": campaigns,
    "recent_attacks": recent_attacks
    }



# =====================================================================================
# ROUTES
# =====================================================================================


@app.route("/", methods=["GET", "POST"])
def index():
    results = {}
    metadata = {}
    if request.method == "POST":
        app_name = request.form.get("app_name", "").strip()
        if app_name:
            # --- Version search ---
            search_query = (
                f"{app_name} latest version release "
                f"site:github.com OR changelog OR release notes"
            )
            try:
                search_results = list(search(search_query, num_results=10))
            except Exception:
                search_results = []

            version, date, source_url = None, None, None
            for url in search_results:
                v, d = extract_version_and_date(url)
                if v:
                    version, date, source_url = v, d, url
                    break

            domain = resolve_application_domain(app_name)

            # --- Description ---
            description, description_url = fetch_app_description(
                app_name, search_results
            )

            # --- CVEs ---
            cve_list = search_vulnerabilities_nvd(app_name)
            if cve_list:
                severity, safe_to_install, breakdown = analyze_risk(cve_list)
                scores = [
                    cve["cvss"]
                    for cve in cve_list
                    if isinstance(cve["cvss"], (int, float))
                ]
                if scores:
                    highest_score = max(scores)
                    avg_score = round(sum(scores) / len(scores), 2)
                else:
                    highest_score, avg_score = "N/A", "N/A"
            else:
                severity = None
                safe_to_install = None
                highest_score = None
                avg_score = None
                breakdown = None

            # --- Security & Compliance ---
            summary, controls, report = {}, {}, None
            if domain != "unknown":
                ports = probe_ports(domain)
                ssl_ok = check_ssl(domain)
                summary = {
                    "domain": domain,
                    "ssl": "valid" if ssl_ok else "invalid",
                    "open_ports": list(ports.keys()),
                }
                controls = {
                    "ISO27001 – A.10.1 (Cryptographic Controls)": "Pass"
                    if ssl_ok
                    else "Fail",
                    "NIST 800-53 – SC-7 (Boundary Protection)": "Fail"
                    if any(p in ports for p in RISKY_PORTS)
                    else "Pass",
                }
                report = llm_report(app_name, domain, summary, controls)
                metadata = fetch_app_metadata(app_name, domain)

            # --- Threat Intelligence ---
            threat_intel = fetch_threat_intel(app_name, domain, cve_list)

            results = {
                "app_name": app_name,
                "version": version,
                "date": date,
                "source_url": source_url,
                "description": description,
                "description_url": description_url,
                "cve_list": cve_list,
                "severity": severity,
                "safe_to_install": safe_to_install,
                "highest_score": highest_score,
                "avg_score": avg_score,
                "breakdown": breakdown,
                "security_summary": summary,
                "compliance": [
                    {"standard": k, "status": v}
                    for k, v in controls.items()
                ],
                "security_analysis": report,
                "threat_intel": threat_intel,
            }

            # Use in-memory cache + lightweight session key
            report_id = str(uuid.uuid4())
            report_cache[report_id] = (results, metadata)
            session["report_id"] = report_id

    return render_template("index.html", results=results, metadata=metadata)


# =====================================================================================
# PDF Report Download
# =====================================================================================


@app.route("/download_report", methods=["POST"])
def download_report():
    app_name = request.form.get("app_name")
    report_id = session.get("report_id")

    if not report_id or report_id not in report_cache:
        return "No report data available. Please analyze an app first.", 400

    results, metadata = report_cache[report_id]

    if not results or results.get("app_name") != app_name:
        return "No report data found for this application.", 400

    # ---- PDF Generation ----
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    x_margin, y_margin = 50, 50
    y = height - y_margin
    max_line_width = width

    def write_line(text, font="Helvetica", size=10, leading=14, bullet=False):
        nonlocal y
        p.setFont(font, size)
        max_chars = int(max_line_width / (size * 0.55))
        lines = textwrap.wrap(text, width=max_chars) or [""]

        for i, line in enumerate(lines):
            if y < y_margin:
                p.showPage()
                y = height - y_margin
                p.setFont(font, size)

            prefix = "• " if bullet and i == 0 else "   " if bullet else ""
            p.drawString(x_margin, y, prefix + line)
            y -= leading

    def write_section(title, size=14):
        nonlocal y
        y -= 10
        p.setFont("Helvetica-Bold", size)
        p.drawString(x_margin, y, title)
        y -= 6
        p.setLineWidth(0.5)
        p.line(x_margin, y, width - x_margin, y)
        y -= 14

    # ---- Header ----
    p.setTitle(f"{app_name} Security Report")
    p.setFont("Helvetica-Bold", 18)
    p.drawCentredString(
        width / 2, y, f"Security & Compliance Report: {app_name}"
    )
    y -= 30
    p.setFont("Helvetica", 10)
    p.drawCentredString(
        width / 2,
        y,
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    )
    y -= 40

    # ---- Metadata ----
    write_section("Application Metadata")
    write_line(f"Version: {metadata.get('latest_version', 'N/A')}")
    write_line(f"Release Date: {metadata.get('release_date', 'N/A')}")
    write_line(f"Support Status: {metadata.get('support_status','N/A')}")
    write_line(f"App Type: {metadata.get('app_type','N/A')}")
    write_line(f"License: {metadata.get('license','N/A')}")
    write_line(f"Official Source: {metadata.get('official_link', 'N/A')}")
    write_line(f"Description URL: {results.get('description_url', 'N/A')}")

    # ---- Description ----
    write_section("Description")
    desc = metadata.get("description") or "N/A"
    for line in str(desc).split("\n"):
        write_line(line.strip())

    # ---- CVE Summary ----
    write_section("CVE Summary")
    write_line(f"Severity: {results.get('severity')}")
    write_line(f"Safe to Install: {results.get('safe_to_install')}")
    write_line(f"Highest CVSS Score: {results.get('highest_score')}")
    write_line(f"Average CVSS Score: {results.get('avg_score')}")
    breakdown = results.get("breakdown", {})
    for k, v in breakdown.items():
        write_line(f"{k.capitalize()}: {v}")

    # ---- CVE List ----
    write_section("Top Vulnerabilities")
    cve_list = results.get("cve_list", [])[:10]
    if not cve_list:
        write_line("No known vulnerabilities found.")
    for cve in cve_list:
        write_line(
            f"{cve['id']}  (CVSS {cve['cvss']})",
            font="Helvetica-Bold",
            bullet=True,
        )
        write_line(f"{cve['summary'][:200]}...", leading=12)

    # ---- Security ----
    write_section("Security Summary")
    sec = results.get("security_summary", {})
    for k, v in sec.items():
        write_line(f"{k}: {v}", bullet=True)

    # ---- Compliance ----
    write_section("Compliance Controls")
    for c in results.get("compliance", []):
        write_line(
            f"{c['standard']}: {c['status']}",
            bullet=True,
        )

    # ---- LLM Report ----
    write_section("LLM Security Analysis")
    report = results.get("security_analysis") or "N/A"
    for line in str(report).split("\n"):
        write_line(line.strip())

    # ---- Threat Intel Summary (short) ----
    ti = results.get("threat_intel", {})
    write_section("Threat Intelligence (Summary)")
    recent = ti.get("recent_attacks", [])


    write_section("Recent Security Attacks / News")
    if recent:
        for a in recent[:10]:
            write_line(a["title"], bullet=True)
            write_line(a["url"])
            write_line(str(a["date"]))
    else:
        write_line("No recent attacks found.")


    p.showPage()
    p.save()
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{app_name}_security_report.pdf",
        mimetype="application/pdf",
    )


# =====================================================================================
# DOCX Report Download
# =====================================================================================


@app.route("/download_report_docx", methods=["POST"])
def download_report_docx():
    app_name = request.form.get("app_name")
    report_id = session.get("report_id")

    if not report_id or report_id not in report_cache:
        return "No report data available. Please analyze an app first.", 400

    results, metadata = report_cache[report_id]

    if not results or results.get("app_name") != app_name:
        return "No report data found for this application.", 400

    doc = Document()

    # Title
    doc.add_heading(f"Security & Compliance Report: {app_name}", level=0)
    doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    # Application Metadata
    doc.add_heading("Application Metadata", level=1)
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = "Light Grid"

    def add_meta_row(label, value):
        row_cells = meta_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value if value is not None else "N/A")

    add_meta_row("Latest Version", metadata.get("latest_version", "N/A"))
    add_meta_row("Release Date", metadata.get("release_date", "N/A"))
    add_meta_row("Support Status", metadata.get("support_status", "N/A"))
    add_meta_row("App Type", metadata.get("app_type", "N/A"))
    add_meta_row("License", metadata.get("license", "N/A"))
    add_meta_row("Official Source", metadata.get("official_link", "N/A"))
    add_meta_row("Description URL", results.get("description_url", "N/A"))

    # Description
    doc.add_heading("Description", level=1)
    desc = metadata.get("description") or "N/A"
    for line in str(desc).split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    # CVE Summary
    doc.add_heading("CVE Summary", level=1)
    doc.add_paragraph(f"Severity: {results.get('severity')}")
    doc.add_paragraph(f"Safe to Install: {results.get('safe_to_install')}")
    doc.add_paragraph(f"Highest CVSS Score: {results.get('highest_score')}")
    doc.add_paragraph(f"Average CVSS Score: {results.get('avg_score')}")

    breakdown = results.get("breakdown", {})
    if breakdown:
        bk_table = doc.add_table(rows=0, cols=2)
        bk_table.style = "Light Grid"
        for k, v in breakdown.items():
            row_cells = bk_table.add_row().cells
            row_cells[0].text = k.capitalize()
            row_cells[1].text = str(v)

    # Top Vulnerabilities
    doc.add_heading("Top Vulnerabilities", level=1)
    cve_list = results.get("cve_list", [])[:10]
    if not cve_list:
        doc.add_paragraph(
            "No known vulnerabilities found in NVD for this software."
        )
    else:
        for cve in cve_list:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{cve['id']} (CVSS {cve['cvss']})").bold = True
            doc.add_paragraph(str(cve["summary"]), style="List Continue")

    # Security Summary
    doc.add_heading("Security Summary", level=1)
    sec = results.get("security_summary", {})
    if sec:
        for k, v in sec.items():
            doc.add_paragraph(f"{k}: {v}", style="List Bullet")
    else:
        doc.add_paragraph("No security summary available.")

    # Compliance
    doc.add_heading("Compliance Controls", level=1)
    compliance = results.get("compliance", [])
    if compliance:
        for c in compliance:
            doc.add_paragraph(
                f"{c['standard']}: {c['status']}", style="List Bullet"
            )
    else:
        doc.add_paragraph("No compliance information available.")

    # Threat Intel
    doc.add_heading("Threat Intelligence", level=1)
    ti = results.get("threat_intel", {}) or {}
    recent = ti.get("recent_attacks", [])
    doc.add_heading("Recent Security Attacks / News", level=2)

    # LLM Security Analysis
    doc.add_heading("LLM Security Analysis", level=1)
    report_text = results.get("security_analysis") or "N/A"
    for line in str(report_text).split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{app_name}_security_report.docx",
        mimetype=(
            "application/"
            "vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


# =====================================================================================
# MAIN
# =====================================================================================

if __name__ == "__main__":
    # app.run(debug=True, host="0.0.0.0", port=int(os.getenv("PORT", 5000)))
    app.run(
        debug=True,
        host="localhost",
        port=int(os.getenv("PORT", 5050)),
    )

